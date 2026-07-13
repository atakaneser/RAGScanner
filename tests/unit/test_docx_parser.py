"""Synthetic, in-memory tests for the bounded DOCX parser."""

from collections.abc import Callable
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from ragscanner.domain import SourceContent, SourceItem
from ragscanner.parsers import (
    BLOCK_SEPARATOR,
    DOCX_MIME,
    DocxParser,
    DocxParserConfig,
    DocxParserError,
    DocxParserErrorCategory,
)

NOW = datetime(2026, 7, 12, 12, 0, tzinfo=UTC)


def make_docx(build: Callable[[DocxDocument], None] | None = None) -> bytes:
    document = Document()
    if build:
        build(document)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def rewrite_zip(
    data: bytes,
    replacements: dict[str, bytes] | None = None,
    additions: dict[str, bytes] | None = None,
) -> bytes:
    replacements = replacements or {}
    output = BytesIO()
    with ZipFile(BytesIO(data)) as source, ZipFile(output, "w", ZIP_DEFLATED) as target:
        for item in source.infolist():
            target.writestr(item, replacements.get(item.filename, source.read(item.filename)))
        for name, value in (additions or {}).items():
            target.writestr(name, value)
    return output.getvalue()


def source(data: bytes, path: str = "synthetic.docx", mime: str = DOCX_MIME) -> SourceContent:
    return SourceContent(
        item=SourceItem(
            id="docx-item",
            source_id="local-filesystem",
            external_id=path,
            name=Path(path).name,
            path=path,
            mime_type=mime,
            size_bytes=len(data),
            modified_at=NOW,
        ),
        content_bytes=data,
        content_type=mime,
        retrieved_at=NOW,
        limit_bytes=max(1, len(data)),
    )


def parse(data: bytes, config: DocxParserConfig | None = None):  # type: ignore[no-untyped-def]
    return DocxParser(config, clock=lambda: NOW, monotonic_clock=lambda: 0).parse(source(data))


def codes(result) -> set[str]:  # type: ignore[no-untyped-def]
    return {warning.code for warning in result.warnings}


def test_simple_document_metadata_title_hash_and_turkish_text() -> None:
    def build(document: DocxDocument) -> None:
        document.core_properties.title = "Kurumsal Bilgi"
        document.core_properties.author = "Test Yazarı"
        document.core_properties.language = "tr-TR"
        document.add_paragraph("Merhaba dünya — hello world")

    data = make_docx(build)
    first = parse(data)
    second = parse(data)
    assert first.document.title == "Kurumsal Bilgi"
    assert first.document.language == "tr-TR"
    assert first.document.content_hash == second.document.content_hash
    assert first.document.metadata["core_properties"]["creator"] == "Test Yazarı"
    assert first.metadata["chunked"] is False


def test_paragraph_heading_lists_table_order_and_offsets() -> None:
    def build(document: DocxDocument) -> None:
        document.add_heading("Başlık", level=1)
        document.add_paragraph("Önce")
        document.add_paragraph("Madde", style="List Bullet")
        document.add_paragraph("İkinci", style="List Number 2")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).merge(table.cell(1, 1)).text = "Birleşik"
        document.add_paragraph("Sonra")

    result = parse(make_docx(build))
    blocks = result.document.metadata["blocks"]
    texts = [block["text"] for block in blocks]
    assert texts.index("Önce") < texts.index("A") < texts.index("Sonra")
    assert any(block["block_type"] == "heading" and block["heading_level"] == 1 for block in blocks)
    assert any(
        block["block_type"] == "list_item" and block["list_ordered"] is False for block in blocks
    )
    assert any(block["block_type"] == "list_item" and block["list_level"] == 1 for block in blocks)
    assert any(
        block["metadata"]["merged"] for block in blocks if block["block_type"] == "table_cell"
    )
    for block in blocks:
        assert result.document.content[block["start_offset"] : block["end_offset"]] == block["text"]


def test_empty_cells_nested_table_and_empty_document_warnings() -> None:
    def build(document: DocxDocument) -> None:
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = ""
        table.cell(0, 1).add_table(rows=1, cols=1).cell(0, 0).text = "nested"

    nested = parse(make_docx(build))
    assert "nested_table_partial_support" in codes(nested)
    assert any(block["text"] == "" for block in nested.document.metadata["blocks"])
    assert "empty_document" in codes(parse(make_docx()))


def test_sections_headers_footers_and_page_breaks() -> None:
    def build(document: DocxDocument) -> None:
        document.sections[0].header.paragraphs[0].text = "Repeated header"
        document.sections[0].footer.paragraphs[0].text = "Footer"
        paragraph = document.add_paragraph("Page one")
        paragraph.add_run().add_break(WD_BREAK.PAGE)
        section = document.add_section(WD_SECTION.NEW_PAGE)
        section.header.is_linked_to_previous = True
        document.add_paragraph("Page two")

    result = parse(make_docx(build))
    kinds = [block["block_type"] for block in result.document.metadata["blocks"]]
    assert {"header", "footer", "page_break", "section_break"}.issubset(kinds)
    assert result.metadata["statistics"]["sections"] == 2


def test_hyperlink_external_relationship_is_recorded_but_never_fetched() -> None:
    def build(document: DocxDocument) -> None:
        paragraph = document.add_paragraph("Visit ")
        relation = paragraph.part.relate_to(
            "https://example.invalid/never-fetch", RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )
        paragraph.part.relate_to(
            "https://example.invalid/external-image.png",
            RELATIONSHIP_TYPE.IMAGE,
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relation)
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "safe label"
        run.append(text)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    result = parse(make_docx(build))
    assert "safe label" in result.document.content
    assert result.document.metadata["hyperlinks"][0]["target"].endswith("never-fetch")
    assert any(
        relationship["type"] == "image"
        for relationship in result.document.metadata["external_relationships"]
    )
    assert "suspicious_external_relationship" in codes(result)
    assert result.metadata["external_resources_fetched"] is False


def test_comments_tracked_changes_hidden_text_and_active_content_warnings() -> None:
    def build(document: DocxDocument) -> None:
        paragraph = document.add_paragraph("Visible ")
        hidden = paragraph.add_run("SECRET HIDDEN")
        hidden.font.hidden = True
        document.add_comment(paragraph.runs[:1], text="review", author="tester")

    data = make_docx(build)
    with ZipFile(BytesIO(data)) as package:
        xml = package.read("word/document.xml")
    xml = xml.replace(
        b"</w:p>",
        b"<w:ins><w:r><w:t>Inserted</w:t></w:r></w:ins>"
        b"<w:del><w:r><w:delText>DELETED</w:delText></w:r></w:del></w:p>",
        1,
    )
    data = rewrite_zip(
        data,
        {"word/document.xml": xml},
        {"word/embeddings/oleObject1.bin": b"never execute", "word/vbaProject.bin": b"macro"},
    )
    result = parse(data)
    assert "Inserted" in result.document.content
    assert "SECRET HIDDEN" not in result.document.content
    assert "DELETED" not in result.document.content
    assert {
        "comments_present",
        "tracked_changes_present",
        "hidden_text_present",
        "embedded_object_present",
        "macro_enabled_document",
    }.issubset(codes(result))
    assert result.document.metadata["active_content_executed"] is False
    assert result.document.metadata["embedded_objects_extracted"] is False


def test_reserved_separator_is_escaped() -> None:
    data = make_docx(lambda document: document.add_paragraph(BLOCK_SEPARATOR))
    result = parse(data)
    blocks = result.document.metadata["blocks"]
    assert BLOCK_SEPARATOR not in blocks[0]["text"]
    assert "block_separator_escaped" in codes(result)


@pytest.mark.parametrize(
    ("data", "path", "category"),
    [
        (b"not a zip", "broken.docx", DocxParserErrorCategory.MALFORMED),
        (b"legacy", "legacy.doc", DocxParserErrorCategory.UNSUPPORTED),
        (b"macro", "macro.docm", DocxParserErrorCategory.UNSUPPORTED),
        (b"\xd0\xcf\x11\xe0garbage", "encrypted.docx", DocxParserErrorCategory.ENCRYPTED),
    ],
)
def test_unsupported_malformed_and_encrypted_inputs(data: bytes, path: str, category) -> None:  # type: ignore[no-untyped-def]
    with pytest.raises(DocxParserError) as caught:
        DocxParser().parse(source(data, path=path))
    assert caught.value.category is category


def test_malformed_xml_fails_safely() -> None:
    data = make_docx(lambda document: document.add_paragraph("valid"))
    data = rewrite_zip(data, {"word/document.xml": b"<broken"})
    with pytest.raises(DocxParserError) as caught:
        parse(data)
    assert caught.value.category is DocxParserErrorCategory.MALFORMED


def test_file_entry_decompressed_compression_and_character_limits() -> None:
    data = make_docx(lambda document: document.add_paragraph("0123456789"))
    many = rewrite_zip(data, additions={f"customXml/{index}.dat": b"x" for index in range(5)})
    compressed = rewrite_zip(data, additions={"word/media/bomb.bin": b"0" * 1_100_000})
    cases = [
        (DocxParserConfig(maximum_file_size=len(data) - 1), data),
        (DocxParserConfig(maximum_zip_entries=5), many),
        (DocxParserConfig(maximum_decompressed_bytes=100), data),
        (DocxParserConfig(maximum_compression_ratio=2), compressed),
        (DocxParserConfig(maximum_extracted_characters=5), data),
    ]
    for config, value in cases:
        with pytest.raises(DocxParserError) as caught:
            parse(value, config)
        assert caught.value.category in {
            DocxParserErrorCategory.LIMIT_EXCEEDED,
            DocxParserErrorCategory.ZIP_BOMB_RISK,
        }


def test_paragraph_table_cell_and_timeout_limits() -> None:
    def build(document: DocxDocument) -> None:
        document.add_paragraph("one")
        document.add_paragraph("two")
        document.add_table(rows=1, cols=2)

    data = make_docx(build)
    limits = [
        DocxParserConfig(maximum_paragraphs=1),
        DocxParserConfig(maximum_tables=1, maximum_cells=1),
    ]
    for config in limits:
        with pytest.raises(DocxParserError) as caught:
            parse(data, config)
        assert caught.value.category is DocxParserErrorCategory.LIMIT_EXCEEDED
    ticks = iter([0.0, 0.0, 2.0])
    parser = DocxParser(
        DocxParserConfig(timeout_seconds=1),
        clock=lambda: NOW,
        monotonic_clock=lambda: next(ticks, 2.0),
    )
    with pytest.raises(DocxParserError) as timeout:
        parser.parse(source(data))
    assert timeout.value.category is DocxParserErrorCategory.TIMEOUT


def test_metadata_is_bounded_and_secret_like_values_are_not_exposed() -> None:
    def build(document: DocxDocument) -> None:
        document.core_properties.title = "x" * 200
        document.core_properties.author = "api_key=super-secret-metadata-value"
        document.add_paragraph("content")

    result = parse(make_docx(build), DocxParserConfig(maximum_metadata_field_length=64))
    metadata = result.document.metadata["core_properties"]
    assert len(metadata["title"]) <= 64
    assert "super-secret" not in str(metadata)


def test_malformed_core_date_is_omitted_with_warning() -> None:
    data = make_docx(lambda document: document.add_paragraph("content"))
    with ZipFile(BytesIO(data)) as package:
        core = package.read("docProps/core.xml")
    core = core.replace(b"2013-12-23T23:15:00Z", b"not-a-date", 1)
    result = parse(rewrite_zip(data, {"docProps/core.xml": core}))
    assert "malformed_metadata" in codes(result)
    assert "created" not in result.document.metadata["core_properties"]


def test_filename_fallback_and_mime_only_acceptance() -> None:
    data = make_docx()
    result = DocxParser(clock=lambda: NOW, monotonic_clock=lambda: 0).parse(
        source(data, path="no-extension", mime=DOCX_MIME)
    )
    assert result.document.title == "no-extension"
