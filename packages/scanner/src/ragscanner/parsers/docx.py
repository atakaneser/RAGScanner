"""Bounded, non-rendering DOCX structure parser using python-docx."""

import re
import zipfile
from collections import Counter
from collections.abc import Callable
from datetime import UTC, datetime
from enum import StrEnum
from io import BytesIO
from pathlib import PurePosixPath
from time import monotonic
from typing import Any
from xml.etree.ElementTree import ParseError

from defusedxml import ElementTree as SafeET
from docx import Document as OpenDocument
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from pydantic import BaseModel, Field

from ragscanner.domain import SourceContent
from ragscanner.domain.helpers import (
    REDACTED,
    contains_unreferenced_secret,
    mask_secret_like_values,
    normalize_control_characters,
    truncate_text,
)
from ragscanner.parsers.base import ParserResult, ParserWarning, build_document

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
BLOCK_SEPARATOR = "\n<<<RAGSCANNER_DOCX_BLOCK_BOUNDARY:4C10A82E>>>\n"
ESCAPED_BLOCK_SEPARATOR = "<<<RAGSCANNER_DOCX_BLOCK_BOUNDARY_ESCAPED:4C10A82E>>>"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
HEADING_STYLE = re.compile(
    r"^(?:heading|başlık|überschrift|titre|título|titolo|nagłówek|заголовок)\s*([1-9])$",
    re.IGNORECASE,
)


class DocxParserErrorCategory(StrEnum):
    UNSUPPORTED = "unsupported"
    MALFORMED = "malformed"
    ENCRYPTED = "encrypted"
    LIMIT_EXCEEDED = "limit_exceeded"
    ZIP_BOMB_RISK = "zip_bomb_risk"
    TIMEOUT = "timeout"


class DocxParserError(Exception):
    def __init__(self, category: DocxParserErrorCategory, message: str) -> None:
        self.category = category
        super().__init__(mask_secret_like_values(message))

    def __repr__(self) -> str:
        return f"DocxParserError(category={self.category.value!r}, message={str(self)!r})"


class DocxParserConfig(BaseModel):
    maximum_file_size: int = Field(default=25 * 1024 * 1024, gt=0)
    maximum_zip_entries: int = Field(default=2_000, gt=0)
    maximum_decompressed_bytes: int = Field(default=100 * 1024 * 1024, gt=0)
    maximum_xml_part_size: int = Field(default=10 * 1024 * 1024, gt=0)
    maximum_compression_ratio: float = Field(default=100, gt=1, le=10_000)
    maximum_paragraphs: int = Field(default=100_000, gt=0)
    maximum_tables: int = Field(default=10_000, gt=0)
    maximum_cells: int = Field(default=500_000, gt=0)
    maximum_extracted_characters: int = Field(default=5_000_000, gt=0)
    maximum_metadata_field_length: int = Field(default=1_024, ge=64, le=16_384)
    timeout_seconds: float = Field(default=30, gt=0, le=600)


class DocxBlockType(StrEnum):
    PARAGRAPH = "paragraph"
    HEADING = "heading"
    LIST_ITEM = "list_item"
    TABLE_CELL = "table_cell"
    HEADER = "header"
    FOOTER = "footer"
    PAGE_BREAK = "page_break"
    SECTION_BREAK = "section_break"


class DocxBlock(BaseModel):
    block_index: int = Field(ge=0)
    block_type: DocxBlockType
    text: str
    start_offset: int = Field(ge=0)
    end_offset: int = Field(ge=0)
    heading_level: int | None = Field(default=None, ge=1, le=9)
    list_ordered: bool | None = None
    list_level: int | None = Field(default=None, ge=0)
    numbering_id: str | None = None
    table_index: int | None = Field(default=None, ge=0)
    row_index: int | None = Field(default=None, ge=0)
    column_index: int | None = Field(default=None, ge=0)
    section_index: int = Field(default=0, ge=0)
    region: str = "body"
    metadata: dict[str, Any] = Field(default_factory=dict)


class DocxStatistics(BaseModel):
    paragraphs: int = 0
    headings: int = 0
    list_items: int = 0
    tables: int = 0
    cells: int = 0
    sections: int = 0
    headers: int = 0
    footers: int = 0
    hyperlinks: int = 0
    embedded_objects_detected: int = 0
    total_characters: int = 0
    warnings_count: int = 0


class PackageInspection(BaseModel):
    warnings: list[ParserWarning] = Field(default_factory=list)
    external_relationships: list[dict[str, str]] = Field(default_factory=list)
    hyperlinks: list[dict[str, str]] = Field(default_factory=list)
    numbering_formats: dict[str, str] = Field(default_factory=dict)
    embedded_objects: int = 0


class DocxParser:
    name = "docx"
    version = "1.0.0"

    def __init__(
        self,
        config: DocxParserConfig | None = None,
        clock: Callable[[], datetime] | None = None,
        monotonic_clock: Callable[[], float] | None = None,
    ) -> None:
        self.config = config or DocxParserConfig()
        self._clock = clock or (lambda: datetime.now(UTC))
        self._monotonic = monotonic_clock or monotonic

    def parse(self, source: SourceContent) -> ParserResult:
        self._validate_source(source)
        if len(source.content_bytes) > self.config.maximum_file_size:
            raise DocxParserError(
                DocxParserErrorCategory.LIMIT_EXCEEDED,
                "DOCX exceeds the configured file-size limit",
            )
        started = self._monotonic()
        inspection = self._inspect_package(source.content_bytes)
        self._check_timeout(started)
        try:
            document = OpenDocument(BytesIO(source.content_bytes))
        except Exception as error:
            raise DocxParserError(
                DocxParserErrorCategory.MALFORMED, "DOCX package or XML is malformed"
            ) from error
        warnings = list(inspection.warnings)
        core = self._core_metadata(document, warnings)
        blocks, stats = self._blocks(document, inspection, warnings, started)
        visible = [block for block in blocks if block.text.strip()]
        if not visible:
            self._warn(warnings, "empty_document", "DOCX has no visible extractable text.")
        elif not any(block.region == "body" for block in visible):
            self._warn(warnings, "header_footer_only", "DOCX contains only header or footer text.")
        content = self._combine(blocks)
        title = (
            core.get("title")
            or next(
                (
                    block.text
                    for block in blocks
                    if block.region == "body" and block.heading_level == 1 and block.text.strip()
                ),
                None,
            )
            or next(
                (block.text for block in blocks if block.region == "body" and block.text.strip()),
                None,
            )
            or PurePosixPath(source.item.path or source.item.name).stem
        )
        stats.sections = len(document.sections)
        stats.hyperlinks = len(inspection.hyperlinks)
        stats.embedded_objects_detected = inspection.embedded_objects
        stats.total_characters = sum(len(block.text) for block in blocks)
        stats.warnings_count = len(warnings)
        parsed = build_document(
            source,
            content=content,
            normalized_content=content,
            title=self._sanitize(str(title), "title"),
            mime_type=DOCX_MIME,
            metadata={
                "core_properties": core,
                "blocks": [block.model_dump(mode="json") for block in blocks],
                "block_separator": BLOCK_SEPARATOR,
                "external_relationships": inspection.external_relationships,
                "hyperlinks": inspection.hyperlinks,
                "source_order_strategy": "document_xml_then_section_headers_and_footers",
                "active_content_executed": False,
                "embedded_objects_extracted": False,
            },
            warnings=warnings,
            clock=self._now(),
            language=str(core["language"]) if core.get("language") else None,
        )
        return ParserResult(
            document=parsed,
            warnings=warnings,
            parser_name=self.name,
            parser_version=self.version,
            source_item_id=source.item.id,
            metadata={
                "statistics": stats.model_dump(mode="json"),
                "chunked": False,
                "active_content_executed": False,
                "embedded_objects_extracted": False,
                "external_resources_fetched": False,
            },
        )

    def _validate_source(self, source: SourceContent) -> None:
        extension = PurePosixPath(source.item.path or source.item.name).suffix.casefold()
        if extension == ".doc":
            raise DocxParserError(DocxParserErrorCategory.UNSUPPORTED, "Legacy DOC is unsupported")
        if extension == ".docm":
            raise DocxParserError(
                DocxParserErrorCategory.UNSUPPORTED, "Macro-enabled DOCM is unsupported"
            )
        if source.content_bytes.startswith(b"\xd0\xcf\x11\xe0"):
            raise DocxParserError(
                DocxParserErrorCategory.ENCRYPTED,
                "Encrypted or legacy OLE Office document is unsupported",
            )
        if (
            source.content_type != DOCX_MIME
            and source.item.mime_type != DOCX_MIME
            and extension != ".docx"
        ):
            raise DocxParserError(
                DocxParserErrorCategory.UNSUPPORTED, "DOCX parser requires DOCX MIME or extension"
            )

    def _inspect_package(self, data: bytes) -> PackageInspection:
        result = PackageInspection()
        try:
            package = zipfile.ZipFile(BytesIO(data))
        except zipfile.BadZipFile as error:
            raise DocxParserError(
                DocxParserErrorCategory.MALFORMED, "DOCX is not a valid ZIP package"
            ) from error
        with package:
            entries = package.infolist()
            if len(entries) > self.config.maximum_zip_entries:
                raise DocxParserError(
                    DocxParserErrorCategory.LIMIT_EXCEEDED, "DOCX exceeds ZIP-entry limit"
                )
            total = 0
            for entry in entries:
                if entry.flag_bits & 1:
                    raise DocxParserError(
                        DocxParserErrorCategory.ENCRYPTED, "Encrypted ZIP entries are unsupported"
                    )
                path = PurePosixPath(entry.filename)
                if path.is_absolute() or ".." in path.parts:
                    raise DocxParserError(
                        DocxParserErrorCategory.MALFORMED, "DOCX contains unsafe ZIP entry path"
                    )
                total += entry.file_size
                if total > self.config.maximum_decompressed_bytes:
                    raise DocxParserError(
                        DocxParserErrorCategory.ZIP_BOMB_RISK,
                        "DOCX decompressed size exceeds safe limit",
                    )
                if (
                    entry.file_size > 1_000_000
                    and entry.compress_size > 0
                    and entry.file_size / entry.compress_size
                    > self.config.maximum_compression_ratio
                ):
                    raise DocxParserError(
                        DocxParserErrorCategory.ZIP_BOMB_RISK,
                        "DOCX compression ratio exceeds safe limit",
                    )
                if (
                    entry.filename.casefold().endswith((".xml", ".rels"))
                    and entry.file_size > self.config.maximum_xml_part_size
                ):
                    raise DocxParserError(
                        DocxParserErrorCategory.LIMIT_EXCEEDED, "DOCX XML part exceeds safe limit"
                    )
                lowered = entry.filename.casefold()
                if "vbaproject.bin" in lowered:
                    self._warn(
                        result.warnings,
                        "macro_enabled_document",
                        "Macro-related content exists; it was not executed.",
                    )
                if lowered.startswith("word/embeddings/") or "oleobject" in lowered:
                    result.embedded_objects += 1
            if result.embedded_objects:
                self._warn(
                    result.warnings,
                    "embedded_object_present",
                    "Embedded objects exist; none were extracted or executed.",
                )
            result.external_relationships = self._relationships(package, result.warnings)
            result.hyperlinks = self._hyperlinks(package, result.external_relationships)
            result.numbering_formats = self._numbering(package)
            self._markup(package, result.warnings)
        return result

    def _relationships(
        self, package: zipfile.ZipFile, warnings: list[ParserWarning]
    ) -> list[dict[str, str]]:
        found: list[dict[str, str]] = []
        for name in sorted(item for item in package.namelist() if item.endswith(".rels")):
            root = self._xml(package, name)
            for relation in root.findall(f"{{{REL_NS}}}Relationship"):
                if relation.attrib.get("TargetMode", "").casefold() != "external":
                    continue
                relation_type = relation.attrib.get("Type", "")
                found.append(
                    {
                        "id": relation.attrib.get("Id", "")[:128],
                        "type": relation_type.rsplit("/", 1)[-1][:128],
                        "target": self._sanitize(relation.attrib.get("Target", ""), "target"),
                        "source_part": name[:256],
                    }
                )
                code = (
                    "external_template_reference"
                    if relation_type.endswith("attachedTemplate")
                    else "suspicious_external_relationship"
                )
                self._warn(warnings, code, "External relationship exists; it was not followed.")
        return found

    def _hyperlinks(
        self, package: zipfile.ZipFile, relationships: list[dict[str, str]]
    ) -> list[dict[str, str]]:
        if "word/document.xml" not in package.namelist():
            return []
        targets = {item["id"]: item["target"] for item in relationships}
        root = self._xml(package, "word/document.xml")
        links: list[dict[str, str]] = []
        for hyperlink in root.findall(f".//{{{W_NS}}}hyperlink"):
            relation_id = hyperlink.attrib.get(f"{{{R_NS}}}id", "")
            text = "".join(node.text or "" for node in hyperlink.findall(f".//{{{W_NS}}}t"))
            links.append(
                {
                    "text": self._sanitize(text, "hyperlink_text"),
                    "target": targets.get(relation_id, ""),
                    "external": str(relation_id in targets).lower(),
                }
            )
        return links

    def _numbering(self, package: zipfile.ZipFile) -> dict[str, str]:
        if "word/numbering.xml" not in package.namelist():
            return {}
        root = self._xml(package, "word/numbering.xml")
        abstract: dict[tuple[str, str], str] = {}
        for definition in root.findall(f"{{{W_NS}}}abstractNum"):
            abstract_id = definition.attrib.get(f"{{{W_NS}}}abstractNumId", "")
            for level in definition.findall(f"{{{W_NS}}}lvl"):
                level_id = level.attrib.get(f"{{{W_NS}}}ilvl", "0")
                fmt = level.find(f"{{{W_NS}}}numFmt")
                if fmt is not None:
                    abstract[(abstract_id, level_id)] = fmt.attrib.get(f"{{{W_NS}}}val", "")
        result: dict[str, str] = {}
        for numbering in root.findall(f"{{{W_NS}}}num"):
            num_id = numbering.attrib.get(f"{{{W_NS}}}numId", "")
            node = numbering.find(f"{{{W_NS}}}abstractNumId")
            abstract_id = "" if node is None else node.attrib.get(f"{{{W_NS}}}val", "")
            for (candidate, level), fmt in abstract.items():
                if candidate == abstract_id:
                    result[f"{num_id}:{level}"] = fmt
        return result

    def _markup(self, package: zipfile.ZipFile, warnings: list[ParserWarning]) -> None:
        names = set(package.namelist())
        if "docProps/core.xml" in names:
            core = self._xml(package, "docProps/core.xml")
            date_namespace = "http://purl.org/dc/terms/"
            for property_name in ("created", "modified"):
                node = core.find(f"{{{date_namespace}}}{property_name}")
                if node is None or not node.text:
                    continue
                try:
                    datetime.fromisoformat(node.text.replace("Z", "+00:00"))
                except ValueError:
                    self._warn(warnings, "malformed_metadata", "A core property is malformed.")
        if "word/comments.xml" in names:
            self._warn(
                warnings,
                "comments_present",
                "Comments exist; comment bodies were not included as visible text.",
            )
        for name in sorted(
            item for item in names if item.startswith("word/") and item.endswith(".xml")
        ):
            root = self._xml(package, name)
            if (
                root.find(f".//{{{W_NS}}}ins") is not None
                or root.find(f".//{{{W_NS}}}del") is not None
            ):
                self._warn(
                    warnings,
                    "tracked_changes_present",
                    "Revision markup exists; deleted text was excluded.",
                )
            if root.find(f".//{{{W_NS}}}vanish") is not None:
                self._warn(
                    warnings,
                    "hidden_text_present",
                    "Hidden text exists; hidden runs were excluded.",
                )

    def _blocks(
        self,
        document: DocxDocument,
        inspection: PackageInspection,
        warnings: list[ParserWarning],
        started: float,
    ) -> tuple[list[DocxBlock], DocxStatistics]:
        blocks: list[DocxBlock] = []
        stats = DocxStatistics()
        section_index = 0
        table_index = 0
        for item in document.iter_inner_content():
            self._check_timeout(started)
            if isinstance(item, Paragraph):
                self._paragraph(blocks, item, section_index, "body", inspection, stats, warnings)
                if self._has_page_break(item):
                    self._append(blocks, DocxBlockType.PAGE_BREAK, "", section_index, "body")
                if item._p.pPr is not None and item._p.pPr.sectPr is not None:
                    self._append(blocks, DocxBlockType.SECTION_BREAK, "", section_index, "body")
                    section_index += 1
            elif isinstance(item, Table):
                self._table(blocks, item, table_index, section_index, "body", stats, warnings)
                table_index += 1
        region_values: list[tuple[str, str]] = []
        for section in document.sections:
            for region_name, region in (("header", section.header), ("footer", section.footer)):
                value = "\n".join(
                    self._visible(paragraph) for paragraph in region.paragraphs
                ).strip()
                if value:
                    region_values.append((region_name, value.casefold()))
        repeated = Counter(region_values)
        for section_index, section in enumerate(document.sections):
            for region_name, region in (("header", section.header), ("footer", section.footer)):
                for item in region.iter_inner_content():
                    if isinstance(item, Paragraph):
                        before = len(blocks)
                        self._paragraph(
                            blocks, item, section_index, region_name, inspection, stats, warnings
                        )
                        if len(blocks) > before:
                            blocks[-1].metadata["repeated_across_sections"] = (
                                repeated[(region_name, blocks[-1].text.casefold())] > 1
                            )
                    elif isinstance(item, Table):
                        self._table(
                            blocks, item, table_index, section_index, region_name, stats, warnings
                        )
                        table_index += 1
        return blocks, stats

    def _paragraph(
        self,
        blocks: list[DocxBlock],
        paragraph: Paragraph,
        section: int,
        region: str,
        inspection: PackageInspection,
        stats: DocxStatistics,
        warnings: list[ParserWarning],
    ) -> None:
        stats.paragraphs += 1
        if stats.paragraphs > self.config.maximum_paragraphs:
            raise DocxParserError(
                DocxParserErrorCategory.LIMIT_EXCEEDED, "DOCX exceeds paragraph limit"
            )
        text = self._visible(paragraph)
        if BLOCK_SEPARATOR in text:
            text = text.replace(BLOCK_SEPARATOR, ESCAPED_BLOCK_SEPARATOR)
            self._warn(
                warnings,
                "block_separator_escaped",
                "Visible text contained reserved block separator.",
            )
        heading = self._heading(paragraph)
        list_info = self._list(paragraph, inspection.numbering_formats)
        if region in {"header", "footer"}:
            kind = DocxBlockType.HEADER if region == "header" else DocxBlockType.FOOTER
            if region == "header":
                stats.headers += 1
            else:
                stats.footers += 1
        elif heading is not None:
            kind = DocxBlockType.HEADING
            stats.headings += 1
        elif list_info:
            kind = DocxBlockType.LIST_ITEM
            stats.list_items += 1
        else:
            kind = DocxBlockType.PARAGRAPH
        self._append(
            blocks,
            kind,
            text,
            section,
            region,
            heading_level=heading,
            list_ordered=list_info[0] if list_info else None,
            list_level=list_info[1] if list_info else None,
            numbering_id=list_info[2] if list_info else None,
            metadata={
                "style": self._sanitize(paragraph.style.name if paragraph.style else "", "style")
            },
        )

    def _table(
        self,
        blocks: list[DocxBlock],
        table: Table,
        table_index: int,
        section: int,
        region: str,
        stats: DocxStatistics,
        warnings: list[ParserWarning],
    ) -> None:
        stats.tables += 1
        if stats.tables > self.config.maximum_tables:
            raise DocxParserError(
                DocxParserErrorCategory.LIMIT_EXCEEDED, "DOCX exceeds table limit"
            )
        repeated_cells = Counter(id(cell._tc) for row in table.rows for cell in row.cells)
        for row_index, row in enumerate(table.rows):
            row_properties = row._tr.trPr
            header_row = (
                row_properties is not None and row_properties.find(qn("w:tblHeader")) is not None
            )
            for column_index, cell in enumerate(row.cells):
                stats.cells += 1
                if stats.cells > self.config.maximum_cells:
                    raise DocxParserError(
                        DocxParserErrorCategory.LIMIT_EXCEEDED, "DOCX exceeds cell limit"
                    )
                if cell.tables:
                    self._warn(
                        warnings,
                        "nested_table_partial_support",
                        "Nested table structure was not recursively emitted.",
                    )
                text = "\n".join(self._visible(paragraph) for paragraph in cell.paragraphs)
                properties = cell._tc.tcPr
                merged = repeated_cells[id(cell._tc)] > 1 or (
                    properties is not None
                    and (properties.gridSpan is not None or properties.vMerge is not None)
                )
                self._append(
                    blocks,
                    DocxBlockType.TABLE_CELL,
                    text,
                    section,
                    region,
                    table_index=table_index,
                    row_index=row_index,
                    column_index=column_index,
                    metadata={"merged": merged, "repeated_header_row": header_row},
                )

    def _append(
        self,
        blocks: list[DocxBlock],
        kind: DocxBlockType,
        text: str,
        section: int,
        region: str,
        **values: Any,
    ) -> None:
        blocks.append(
            DocxBlock(
                block_index=len(blocks),
                block_type=kind,
                text=text,
                start_offset=0,
                end_offset=0,
                section_index=section,
                region=region,
                **values,
            )
        )

    def _combine(self, blocks: list[DocxBlock]) -> str:
        parts: list[str] = []
        offset = 0
        characters = 0
        for index, block in enumerate(blocks):
            if index:
                parts.append(BLOCK_SEPARATOR)
                offset += len(BLOCK_SEPARATOR)
            block.start_offset = offset
            parts.append(block.text)
            offset += len(block.text)
            characters += len(block.text)
            block.end_offset = offset
            if characters > self.config.maximum_extracted_characters:
                raise DocxParserError(
                    DocxParserErrorCategory.LIMIT_EXCEEDED, "DOCX exceeds extracted-character limit"
                )
        return "".join(parts)

    def _visible(self, paragraph: Paragraph) -> str:
        parts: list[str] = []

        def walk(element: Any, hidden: bool = False, deleted: bool = False) -> None:
            deleted = deleted or element.tag == qn("w:del")
            if element.tag == qn("w:r"):
                properties = element.find(qn("w:rPr"))
                hidden = hidden or (
                    properties is not None and properties.find(qn("w:vanish")) is not None
                )
            if not hidden and not deleted:
                if element.tag == qn("w:t"):
                    parts.append(element.text or "")
                elif element.tag == qn("w:tab"):
                    parts.append("\t")
                elif element.tag in {qn("w:br"), qn("w:cr")}:
                    parts.append("\n")
            for child in element:
                walk(child, hidden, deleted)

        walk(paragraph._p)
        return "".join(parts)

    @staticmethod
    def _heading(paragraph: Paragraph) -> int | None:
        match = HEADING_STYLE.match((paragraph.style.name if paragraph.style else "").strip())
        if match:
            return int(match.group(1))
        properties = paragraph._p.pPr
        outline = None if properties is None else properties.find(qn("w:outlineLvl"))
        raw = None if outline is None else outline.get(qn("w:val"))
        return int(raw) + 1 if raw and raw.isdigit() and int(raw) < 9 else None

    @staticmethod
    def _list(paragraph: Paragraph, formats: dict[str, str]) -> tuple[bool | None, int, str] | None:
        properties = paragraph._p.pPr
        numbering = None if properties is None else properties.find(qn("w:numPr"))
        num_id_node = None if numbering is None else numbering.find(qn("w:numId"))
        if num_id_node is None:
            style = (paragraph.style.name if paragraph.style else "").casefold()
            match = re.search(r"(?:list|liste)\s+(bullet|number|madde|numara)(?:\s+(\d+))?", style)
            if match is None:
                return None
            ordered = match.group(1) in {"number", "numara"}
            level = max(0, int(match.group(2) or "1") - 1)
            return ordered, level, "style"
        num_id = num_id_node.get(qn("w:val"), "")
        level_node = numbering.find(qn("w:ilvl")) if numbering is not None else None
        level_value = None if level_node is None else level_node.get(qn("w:val"))
        level = int(level_value) if level_value and level_value.isdigit() else 0
        fmt = formats.get(f"{num_id}:{level}")
        return (None if fmt is None else fmt != "bullet", level, num_id)

    @staticmethod
    def _has_page_break(paragraph: Paragraph) -> bool:
        return any(node.get(qn("w:type")) == "page" for node in paragraph._p.iter(qn("w:br")))

    def _core_metadata(
        self, document: DocxDocument, warnings: list[ParserWarning]
    ) -> dict[str, Any]:
        properties = document.core_properties
        result: dict[str, Any] = {}
        for source_name, target_name in (
            ("title", "title"),
            ("subject", "subject"),
            ("author", "creator"),
            ("last_modified_by", "last_modified_by"),
            ("created", "created"),
            ("modified", "modified"),
            ("category", "category"),
            ("keywords", "keywords"),
            ("comments", "description"),
            ("revision", "revision"),
            ("language", "language"),
        ):
            try:
                value = getattr(properties, source_name)
            except (ValueError, TypeError):
                self._warn(warnings, "malformed_metadata", "A core property is malformed.")
                continue
            if value in (None, ""):
                continue
            result[target_name] = (
                value.replace(tzinfo=UTC).isoformat()
                if isinstance(value, datetime)
                else self._sanitize(str(value), target_name)
            )
        return result

    def _xml(self, package: zipfile.ZipFile, name: str) -> Any:
        try:
            return SafeET.fromstring(package.read(name))
        except (KeyError, OSError, ParseError) as error:
            raise DocxParserError(
                DocxParserErrorCategory.MALFORMED, "DOCX contains malformed XML"
            ) from error

    def _sanitize(self, value: str, key: str) -> str:
        bounded = truncate_text(
            normalize_control_characters(value), self.config.maximum_metadata_field_length
        )
        return (
            REDACTED
            if contains_unreferenced_secret(bounded, parent_key=key)
            else mask_secret_like_values(bounded)
        )

    def _check_timeout(self, started: float) -> None:
        if self._monotonic() - started >= self.config.timeout_seconds:
            raise DocxParserError(
                DocxParserErrorCategory.TIMEOUT, "DOCX parsing exceeded the cooperative timeout"
            )

    def _now(self) -> datetime:
        value = self._clock()
        if value.tzinfo is None:
            raise ValueError("DOCX parser clock must be timezone-aware")
        return value

    @staticmethod
    def _warn(warnings: list[ParserWarning], code: str, message: str) -> None:
        if not any(warning.code == code for warning in warnings):
            warnings.append(ParserWarning(code=code, message=message))
